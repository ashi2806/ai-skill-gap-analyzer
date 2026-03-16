import streamlit as st
import json
import ast
import re
import pandas as pd
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from tenacity import retry, stop_after_attempt, wait_exponential
from config import get_llm_client, get_llm_model

def add_hyperlink(paragraph, url, text):
    """
    A function that places a clickable hyperlink within a paragraph object.
    """
    # This gets access to the document.xml.rels file and adds a new relation id 
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )

    # Create a w:r element
    new_run = OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = OxmlElement('w:rPr')

    # Add color and underline
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run_text = OxmlElement('w:t')
    new_run_text.text = text
    new_run.append(new_run_text)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

@st.cache_data
@retry(stop=stop_after_attempt(2), wait=wait_exponential(multiplier=1, min=2, max=8))
def ai_extract_basic_info(resume_text):
    """AI extracts + cleans contact info"""
    client = get_llm_client()
    
    prompt = f"""Extract ONLY name, email, phone from resume as VALID JSON:
{{"name": "John Doe", "email": "john@email.com", "phone": "+919876543210"}}
Resume: {resume_text[:2000]}"""
    
    model_name = get_llm_model()
    response = client.chat.completions.create(
        model=model_name,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1
    )
    content = response.choices[0].message.content.strip()
    
    # Robustly find JSON structure
    json_match = re.search(r'\{.*\}', content, re.DOTALL)
    if json_match:
        try:
            return json.loads(json_match.group(0))
        except:
            pass
            
    # Fallback attempt: strip markdown
    content = re.sub(r'```json|```', '', content).strip()
    try:
        return json.loads(content)
    except:
        return {"name": "Candidate", "email": "", "phone": ""}

@retry(stop=stop_after_attempt(2), wait=wait_exponential(multiplier=1, min=2, max=8))
def ai_generate_career_plan(missing_skills):
    """AI personalized 30-day plans"""
    client = get_llm_client()
    
    skills_str = ", ".join(missing_skills[:3])
    prompt = f"""Create 3 detailed 30-day learning plans for skills: {skills_str}
Return as Python list with 3 strings only:
["Plan 1...", "Plan 2...", "Plan 3..."]"""
    
    response = client.chat.completions.create(
        model="tinyllama",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3
    )
    content = response.choices[0].message.content.strip()
    content = re.sub(r'```python|```', '', content).strip()
    return ast.literal_eval(content)[:3]

def add_links_to_paragraph(paragraph, text):
    """Detects URLs and adds them as clickable hyperlinks to a paragraph."""
    # Pattern for URLs
    url_pattern = r"(https?://[^\s,]+|www\.[^\s,]+|linkedin\.com/[^\s,]+|github\.com/[^\s,]+)"
    
    parts = re.split(url_pattern, text)
    for part in parts:
        if re.match(url_pattern, part):
            url = part if part.startswith("http") else "https://" + part
            add_hyperlink(paragraph, url, part)
        else:
            paragraph.add_run(part)

def ai_enhance_resume_text(resume_text, missing_skills):
    """
    Absolute Verbatim: Uses Python logic to insert missing skills directly below 
    the existing Skills header, ensuring zero characters of original text are changed.
    """
    if not missing_skills:
        return resume_text
        
    skills_to_add = ", ".join(missing_skills)
    
    # Highly specific but flexible Skill header detection
    # Matches: "Skills", "Skills:", "Technical Skills", etc. on their own line
    pattern = r"(?i)^.*(?:technical\s*)?skills(?:(?:\s*&)?\s*(?:certifications|expertise|competencies))?:?\s*$"
    
    lines = resume_text.splitlines()
    new_lines = []
    found = False
    
    for line in lines:
        new_lines.append(line)
        if not found and re.match(pattern, line.strip(), re.IGNORECASE):
            # Insert the new skills right below the header
            # We add it as a new line to preserve the original header line exactly
            new_lines.append(f"Additional Skills: {skills_to_add}")
            found = True
            
    if found:
        return "\n".join(new_lines)
        
    # Fallback: If no Skills section exists, the user requested to return UNCHANGED
    return resume_text

def convert_text_to_docx(resume_text_str):
    """Converts text to DOCX with professional section borders and alignment."""
    doc = Document()
    
    # helper for horizontal line
    def add_horizontal_line(p):
        p_pr = p._element.get_or_add_pPr()
        p_pbdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6') # 0.75 pt
        bottom.set(qn('w:space'), '1')
        p_pbdr.append(bottom)
        p_pr.append(p_pbdr)

    # Standard Professional Margins
    for section in doc.sections:
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.6)

    # Professional Font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    paragraphs = resume_text_str.split('\n')
    
    # Common Section Headers
    section_patterns = [
        r'^(EDUCATION|EXPERIENCE|WORK\s+EXPERIENCE|PROJECTS|TECHNICAL\s+SKILLS|SKILLS|CERTIFICATIONS|ACHIEVEMENTS|LANGUAGES|SUMMARY|OBJECTIVE):?\s*$',
        r'^[A-Z\s]{5,20}$' # Fallback for all-caps short lines
    ]

    for line in paragraphs:
        text = line.strip()
        if not text:
            doc.add_paragraph()
            continue
            
        p = doc.add_paragraph()
        
        # Section Header Detection
        is_header = False
        for pattern in section_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                is_header = True
                break
        
        # Extra check: if it's very short and capitalized, it's likely a header
        if is_header and len(text) < 30:
            run = p.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            add_horizontal_line(p)
            continue

        # Bullet Points
        if text.startswith('•') or text.startswith('-') or text.startswith('*'):
            clean_text = text.lstrip('•-* ').strip()
            p.style = 'List Bullet'
            p.paragraph_format.left_indent = Inches(0.25)
            add_links_to_paragraph(p, clean_text)
        else:
            # Standard Text
            add_links_to_paragraph(p, text)
            
        p.paragraph_format.space_after = Pt(3)
            
    return doc


def ai_build_template_resume(basic_info, final_skills):
    """Builds a fresh resume using the high-fidelity format.png style"""
    doc = Document()
    
    # Margins (0.5 inch for a clean modern look)
    sections = doc.sections
    for s in sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(0.5)
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    def add_section_header(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(12)
        
        # Professional Horizontal Line
        p_pr = p._element.get_or_add_pPr()
        p_pbdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        p_pbdr.append(bottom)
        p_pr.append(p_pbdr)

    # ---------- Header ----------
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.paragraph_format.space_after = Pt(12)
    run_name = h.add_run(basic_info["name"])
    run_name.bold = True
    run_name.font.size = Pt(20)

    # ---------- Contact ----------
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_str = f"{basic_info.get('email', '')} | {basic_info.get('phone', '')} | Location, City"
    run_contact = c.add_run(contact_str)
    run_contact.font.size = Pt(10)

    # ---------- Skills ----------
    add_section_header("SKILLS")
    skills_para = doc.add_paragraph()
    skills_para.add_run("Technical Skills: ").bold = True
    skills_para.add_run(", ".join(final_skills))

    # ---------- Experience ----------
    add_section_header("EXPERIENCE")
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    
    # Hide Table Borders
    tbl = table._tbl
    tblPr = tbl.xpath('w:tblPr')[0]
    tblBorders = OxmlElement('w:tblBorders')
    for bdr in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        node = OxmlElement(f'w:{bdr}')
        node.set(qn('w:val'), 'none')
        tblBorders.append(node)
    tblPr.append(tblBorders)

    table.columns[0].width = Inches(5.5)
    table.columns[1].width = Inches(2.0)
    
    cells = table.rows[0].cells
    cells[0].paragraphs[0].add_run("Company Name").bold = True
    cells[1].paragraphs[0].add_run("Month Year – Present").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    p = doc.add_paragraph("Role Title")
    p.paragraph_format.left_indent = Pt(10)
    
    doc.add_paragraph("Led XYZ project and improved performance by 20%", style='List Bullet')

    # ---------- Education ----------
    add_section_header("EDUCATION")
    table2 = doc.add_table(rows=1, cols=2)
    # (Hiding borders for table 2 would be similar, skipping for brevity in template unless asked)
    cells2 = table2.rows[0].cells
    cells2[0].paragraphs[0].add_run("University Name").bold = True
    cells2[1].paragraphs[0].add_run("Graduation Year").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    return doc
